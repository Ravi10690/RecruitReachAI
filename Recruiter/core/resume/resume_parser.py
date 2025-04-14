"""
Resume Parser for RecruitReach2.

This module provides functionality for loading and parsing resume files.
"""

import os
import io
from typing import Optional, Union, Dict, Any

from langchain_community.document_loaders import PyPDFLoader
from docx import Document

from Recruiter.utils.file_utils.path_manager import PathManager


class ResumeParser:
    """
    Parser for resume files.
    
    This class provides methods for loading and parsing resume files in
    various formats such as PDF and DOCX.
    """
    
    def __init__(self):
        """Initialize the resume parser."""
        self.path_manager = PathManager()
    
    def load_default_resume(self) -> str:
        """
        Load the default resume file.
        
        Returns:
            Content of the default resume as text.
            
        Raises:
            FileNotFoundError: If the default resume file is not found.
        """
        resume_path = self.path_manager.get_resume_path()
        
        if not os.path.isfile(resume_path):
            raise FileNotFoundError(f"Default resume not found at: {resume_path}")
        
        return self.load_resume_from_file(resume_path)
    
    def load_resume_from_file(self, file_path: str) -> str:
        """
        Load a resume from a file.
        
        Args:
            file_path: Path to the resume file.
            
        Returns:
            Content of the resume as text.
            
        Raises:
            FileNotFoundError: If the resume file is not found.
            ValueError: If the file format is not supported.
        """
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"Resume file not found at: {file_path}")
        
        # Determine file type based on extension
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._parse_pdf(file_path)
        elif file_extension == '.docx':
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def load_resume_from_bytes(
        self,
        file_content: bytes,
        file_name: str
    ) -> str:
        """
        Load a resume from bytes.
        
        Args:
            file_content: Content of the resume file as bytes.
            file_name: Name of the resume file.
            
        Returns:
            Content of the resume as text.
            
        Raises:
            ValueError: If the file format is not supported.
        """
        # Determine file type based on extension
        file_extension = os.path.splitext(file_name)[1].lower()
        
        if file_extension == '.pdf':
            return self._parse_pdf_bytes(file_content)
        elif file_extension == '.docx':
            return self._parse_docx_bytes(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _parse_pdf(self, file_path: str) -> str:
        """
        Parse a PDF file.
        
        Args:
            file_path: Path to the PDF file.
            
        Returns:
            Content of the PDF as text.
        """
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        return '\n'.join([doc.page_content for doc in docs])
    
    def _parse_pdf_bytes(self, file_content: bytes) -> str:
        """
        Parse a PDF file from bytes.
        
        Args:
            file_content: Content of the PDF file as bytes.
            
        Returns:
            Content of the PDF as text.
        """
        from PyPDF2 import PdfReader
        
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)
        text = ""
        
        for page in pdf_reader.pages:
            text += page.extract_text()
        
        return text
    
    def _parse_docx(self, file_path: str) -> str:
        """
        Parse a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            Content of the DOCX as text.
        """
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    
    def _parse_docx_bytes(self, file_content: bytes) -> str:
        """
        Parse a DOCX file from bytes.
        
        Args:
            file_content: Content of the DOCX file as bytes.
            
        Returns:
            Content of the DOCX as text.
        """
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
